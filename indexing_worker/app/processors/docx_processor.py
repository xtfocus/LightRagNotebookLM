"""
DOCX Processor

This module handles text extraction from DOCX files.
"""

import io
from typing import Any
from loguru import logger
import docx

from .base import TextProcessor
from app.core.config import settings


class DOCXProcessor(TextProcessor):
    """Handles DOCX document processing."""
    
    def can_handle(self, source_type: str, mime_type: str = None) -> bool:
        """Check if this processor can handle the given source type and MIME type."""
        if source_type != "document":
            return False
        
        if mime_type is None:
            return True
        
        supported_mime_types = [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]
        
        return mime_type in supported_mime_types
    
    async def extract_text(self, file_data: bytes) -> str:
        """
        Extract text from DOCX file data.
        
        Args:
            file_data: Raw DOCX file data as bytes
            
        Returns:
            Extracted text as string
            
        Raises:
            ValueError: If text extraction fails
        """
        if not isinstance(file_data, bytes):
            raise ValueError("File data must be bytes")
        
        # Validate DOCX signature (ZIP file)
        if not self._is_valid_docx(file_data):
            raise ValueError("Invalid DOCX file format")
        
        try:
            return await self._extract_text_from_docx(file_data)
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"DOCX text extraction failed: {str(e)}")
    
    def _is_valid_docx(self, file_data: bytes) -> bool:
        """
        Check if the file data represents a valid DOCX.
        
        Args:
            file_data: Raw file data
            
        Returns:
            True if the file is a valid DOCX
        """
        # DOCX files are ZIP files, so they start with PK signature
        return file_data.startswith(b'PK')
    
    async def _extract_text_from_docx(self, file_data: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc_file = io.BytesIO(file_data)
            doc = docx.Document(doc_file)
            
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += paragraph.text + "\n"
            
            if not text.strip():
                logger.warning("No text extracted from DOCX - file may be empty or corrupted")
                return ""
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"DOCX text extraction failed: {str(e)}")
    
    def validate_input(self, data: Any) -> None:
        """
        Validate DOCX input data.
        
        Args:
            data: DOCX file data to validate
            
        Raises:
            ValueError: If data is invalid
        """
        super().validate_input(data)
        
        if not isinstance(data, bytes):
            raise ValueError("DOCX data must be bytes")
        
        # Check for minimum DOCX size (very small files are likely invalid)
        if len(data) < settings.MIN_FILE_SIZE_BYTES:
            raise ValueError(f"DOCX file appears to be too small (min {settings.MIN_FILE_SIZE_BYTES} bytes)")
        
        # Check for maximum DOCX size (configurable limit)
        if len(data) > settings.MAX_DOCX_SIZE_BYTES:
            max_size_mb = settings.MAX_DOCX_SIZE_BYTES / (1024 * 1024)
            raise ValueError(f"DOCX file is too large (max {max_size_mb:.0f}MB)") 